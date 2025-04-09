import streamlit as st
import os
import fitz  # PyMuPDF
import google.generativeai as genai
from docx import Document
import shutil
import zipfile
from io import BytesIO
import tempfile

# Configuração do Gemini (substitua pelo seu API key e modelo)
GOOGLE_API_KEY = "AIzaSyBrRaTvJ6wo2-PQUxOwRMQzr1S8KucT79A"
genai.configure(api_key=GOOGLE_API_KEY)
model = genai.GenerativeModel('gemini-1.5-flash-latest')

# Função para extrair texto de um PDF
def extract_text_from_pdf(pdf_path):
    document = fitz.open(pdf_path)
    text = ""
    for page_num in range(len(document)):
        page = document.load_page(page_num)
        text += page.get_text()
    document.close()
    return text

# Função para gerar o relatório usando Gemini
def generate_summary(text):
    question = (
        f"PROMPT PARA GERAÇÃO DE RELATÓRIO JURÍDICO COMPLETO E ABRANGENTE\n\n"
        f"Instruções:\n\n"
        
        f"1. Estrutura e Linguagem:\n"
        f"- Utilize uma linguagem formal, clara e objetiva, com imparcialidade.\n"
        f"- O relatório deve seguir uma ordem cronológica e lógica, destacando os pontos mais relevantes das peças processuais, documentos, decisões e demais atos processuais.\n"
        f"- Em todas as peças analisadas, mencione expressamente os principais argumentos jurídicos, dispositivos legais e precedentes/julgados citados pelas partes ou pelo juízo.\n"
        f"- Descreva de forma cuidadosa e minuciosa todos os documentos anexados ao processo, explicando sua relevância para as alegações das partes.\n\n"
        f"- Nas respostas NÃO USE MARKDOWN em nenhum momento\n\n"

        f"2. Flexibilidade do Relatório:\n"
        f"- Este relatório é adaptável para qualquer peça ou fase processual: petições iniciais, contestações, réplicas, decisões interlocutórias, audiências, sentença, recursos e decisões de instâncias superiores.\n"
        f"- Sempre que uma dessas peças ou fases estiver presente no processo, utilize os tópicos correspondentes. Quando inexistentes, omita os tópicos ou ajuste conforme necessário.\n\n"
        
        f"3. Conteúdo do Relatório:\n"
        f"- Inicie sempre identificando as partes, comarca, valor da causa e magistrado. Em seguida, resuma as alegações e documentos apresentados por cada parte, decisões proferidas e atos processuais subsequentes, conforme as instruções.\n"
        f"- Mencione com destaque os dispositivos legais, precedentes judiciais e julgados relevantes citados em cada peça analisada.\n\n"

        f"Estrutura do Relatório:\n\n"
        
        f"RELATÓRIO\n\n"
        f"Trata-se de [TIPO DE AÇÃO EM MAIÚSCULAS SEM ASPAS] proposta por [NOME DA(S) PARTE(S) AUTORA(S) EM MAIÚSCULAS] contra [NOME DA(S) PARTE(S) REQUERIDA(S) EM MAIÚSCULAS], "
        f"na comarca de [COMARCA EM MAIÚSCULAS], tendo como magistrado [NOME DO MAGISTRADO] e valor da causa fixado em [VALOR DA CAUSA]. "
        f"O objetivo desta demanda é [RESUMIR O PEDIDO CENTRAL – EXEMPLO: obter indenização por danos morais e materiais, pleitear rescisão contratual, etc.].\n\n"
        
        f"I - Alegações da Parte Autora:\n"
        f"Alega a parte autora que:\n\n"
        f"[INSERIR FATOS ALEGADOS PELA PARTE AUTORA]. Por exemplo: “A parte autora narra que adquiriu o produto X da parte ré e que este apresentou defeito irreparável dentro do prazo de garantia.”\n\n"
        
        f"Além disso, apresentou os seguintes documentos para fundamentar suas alegações, destacando a relevância de cada um para sua argumentação:\n\n"
        f"[LISTAR DOCUMENTOS APRESENTADOS PELA PARTE AUTORA]. Exemplo: “Nota fiscal de compra (doc. 1), comprovando a aquisição do produto na data informada; "
        f"e-mails trocados com a assistência técnica (doc. 2), mostrando as tentativas de solução; laudo técnico (doc. 3), que atesta o defeito irreparável.”\n\n"
        
        f"Em suas palavras, a parte autora [TRANSCREVER TRECHOS RELEVANTES DA PETIÇÃO INICIAL].\n\n"
        f"Para reforçar suas alegações, a parte autora argumenta que:\n\n"
        f"[INSERIR ARGUMENTOS JURÍDICOS UTILIZADOS PELA PARTE AUTORA]. Exemplo: “A autora fundamenta sua pretensão no art. 12 do Código de Defesa do Consumidor, alegando que a fabricante é objetivamente responsável pelos defeitos apresentados no produto.”\n"
        
        f"Sustenta ainda que:\n\n"
        f"[OUTRAS CONSIDERAÇÕES DA PARTE AUTORA]. Exemplo: “A autora alega que, mesmo após inúmeras tentativas de contato, a ré se recusou a realizar a troca ou o conserto adequado do produto.”\n\n"
        f"Por fim, requer:\n\n"
        f"[DESCREVER O PEDIDO CENTRAL DA AÇÃO]. Exemplo: “A autora requer a substituição do produto defeituoso ou, alternativamente, o reembolso integral do valor pago, além de indenização por danos morais no valor de R$ 50.000,00.”\n\n"
        
        f"II - Contestação da Parte Requerida:\n"
        f"Em contestação, [NOME DA PARTE REQUERIDA EM MAIÚSCULAS] alegou que:\n\n"
        f"[DESCREVER ALEGAÇÕES PRELIMINARES, SE HOUVER]. Exemplo: “A parte ré arguiu preliminarmente a ilegitimidade passiva, afirmando que não é responsável pela fabricação do produto, sendo apenas distribuidora.”\n\n"
        f"Quanto ao mérito, a parte ré sustenta que:\n\n"
        f"[RESUMO DOS FATOS CONTRAPOSTOS PELA PARTE REQUERIDA]. Exemplo: “No mérito, a ré nega a existência de qualquer defeito de fabricação, sustentando que o problema decorre de mau uso por parte da autora.”\n\n"
        
        f"A parte ré anexou os seguintes documentos para comprovar suas alegações, explicando detalhadamente sua importância:\n\n"
        f"[LISTAR DOCUMENTOS APRESENTADOS PELA PARTE REQUERIDA]. Exemplo: “Manual do produto (doc. 1), que descreve o uso adequado; laudo técnico próprio (doc. 2), concluindo que não há defeito de fabricação; "
        f"registro de atendimentos realizados (doc. 3), demonstrando que a ré atendeu tempestivamente às solicitações da autora.”\n\n"
        
        f"Argumenta ainda que:\n\n"
        f"[ARGUMENTOS JURÍDICOS UTILIZADOS NA CONTESTAÇÃO]. Exemplo: “A parte ré sustenta que, conforme o art. 18 do Código de Defesa do Consumidor, foi respeitado o prazo de 30 dias para reparação do vício, e que, portanto, não há razão para a substituição do produto.”\n"
        
        f"Sustenta ainda que:\n\n"
        f"[OUTRAS CONSIDERAÇÕES RELEVANTES DA CONTESTAÇÃO]. Exemplo: “A ré também argumenta que o valor solicitado pela autora a título de danos morais é desproporcional ao suposto prejuízo sofrido.”\n\n"
        f"Por fim, requer:\n\n"
        f"[DESCREVER O PEDIDO CENTRAL DA CONTESTAÇÃO]. Exemplo: “A ré requer a total improcedência dos pedidos formulados pela autora e, subsidiariamente, a redução do valor pleiteado a título de danos morais.”\n\n"
        
        f"III - Réplica da Parte Autora (se aplicável):\n"
        f"Em réplica, a parte autora [NOME DA PARTE AUTORA EM MAIÚSCULAS] refutou as alegações da parte ré, sustentando que:\n\n"
        f"[DESCREVER A RESPOSTA DA PARTE AUTORA À CONTESTAÇÃO]. Exemplo: “A parte autora reafirma que o defeito apresentado no produto decorre de falha de fabricação, como comprovado pelo laudo técnico juntado na inicial.”\n\n"
        
        f"A parte autora apresentou os seguintes documentos em réplica, explicando sua relevância:\n\n"
        f"[LISTAR DOCUMENTOS APRESENTADOS NA RÉPLICA, SE HOUVER].\n\n"
        f"Por fim, a parte autora requer que:\n\n"
        f"[DESCREVER O PEDIDO FINAL APÓS A RÉPLICA]. Exemplo: “A parte autora reitera os pedidos formulados na inicial.”\n\n"
        
        f"IV - Decisões Interlocutórias (se houver):\n"
        f"O magistrado proferiu decisão interlocutória determinando que:\n\n"
        f"[DESCREVER A DECISÃO INTERLOCUTÓRIA]. Exemplo: “O juiz determinou a realização de perícia técnica no produto para verificar a alegada existência de defeito de fabricação, nomeando o perito e fixando prazo para entrega do laudo.”\n\n"
        
        f"V - Audiências (se houver):\n"
        f"Realizou-se audiência de [tipo da audiência, exemplo: conciliação, instrução e julgamento] em [data da audiência]. Na ocasião:\n\n"
        f"[RESUMIR OS EVENTOS DA AUDIÊNCIA]. Exemplo: “As partes não chegaram a acordo durante a audiência de conciliação, sendo designada audiência de instrução para oitiva de testemunhas.”\n\n"
        
        f"VI - Sentença (se houver):\n"
        f"Na sentença, o magistrado julgou que:\n\n"
        f"[RESUMIR O TEOR DA SENTENÇA]. Exemplo: “O juiz julgou parcialmente procedente a ação, condenando a ré a substituir o produto defeituoso e ao pagamento de R$ 10.000,00 por danos morais.”\n\n"
        
        f"VII - Recursos:\n\n"
        f"1. Agravo de Instrumento (se houver):\n"
        f"[NOME DA PARTE QUE INTERPOS AGRAVO] interpôs agravo de instrumento contra a decisão interlocutória que:\n\n"
        f"[DESCREVER A DECISÃO AGRAVADA E OS FUNDAMENTOS DO AGRAVO].\n\n"
        f"Em contraminuta, [NOME DA PARTE CONTRAQUEM FOI INTERPOSTO O AGRAVO] alegou que:\n\n"
        f"[DESCREVER A RESPOSTA DA PARTE EM CONTRAMINUTA].\n\n"
        
        f"2. Apelação (se houver):\n"
        f"[NOME DA PARTE QUE APELOU] interpôs apelação contra a sentença, argumentando que:\n\n"
        f"[DESCREVER OS FUNDAMENTOS DA APELAÇÃO].\n\n"
        f"Em contrarrazões, [NOME DA PARTE QUE APRESENTOU CONTRARRAZÕES] sustentou que:\n\n"
        f"[DESCREVER A RESPOSTA DA PARTE EM CONTRARRAZÕES].\n\n"
        
        f"VIII - Decisão Monocrática de Recurso (se houver):\n"
        f"O relator do recurso decidiu monocraticamente que:\n\n"
        f"[DESCREVER O TEOR DA DECISÃO MONOCRÁTICA].\n\n"
        
        f"IX - Decisão Definitiva de Recurso (se houver):\n"
        f"O tribunal, em decisão colegiada, julgou que:\n\n"
        f"[RESUMIR A DECISÃO FINAL DO RECURSO]. Exemplo: “O tribunal negou provimento à apelação da ré, mantendo integralmente a sentença de primeira instância.”\n\n"
        
        f"X -  Sugestão de Minuta para Capítulo dos Fatos:\n"
        f"Abaixo está um modelo pré-estruturado para a parte 'Dos Fatos', que deve ser adaptado com as informações específicas do caso:\n\n"
        
       f"DOS FATOS\n\n"
        f"Trata-se de [TIPO DE AÇÃO EM MAIÚSCULAS SEM ASPAS] proposta por [NOME DA(S) PARTE(S) AUTORA(S) EM MAIÚSCULAS] contra [NOME DA(S) PARTE(S) REQUERIDA(S) EM MAIÚSCULAS], "
        f"na comarca de [COMARCA EM MAIÚSCULAS], tendo como magistrado [NOME DO MAGISTRADO] e valor da causa fixado em [VALOR DA CAUSA]. "
        f"O objetivo desta demanda é [RESUMIR O PEDIDO CENTRAL – EXEMPLO: obter indenização por danos morais e materiais, pleitear rescisão contratual, etc.].\n\n"
        
        f"Na petição inicial, o autor argumentou que [DESCREVER OS PRINCIPAIS FATOS ALEGADOS PELO AUTOR]. "
        f"Para fundamentar suas alegações, apresentou os seguintes documentos: [LISTAR DOCUMENTOS APRESENTADOS PELO AUTOR].\n\n"
        
        f"Em contestação, a parte ré alegou [DESCREVER OS PRINCIPAIS ARGUMENTOS DO RÉU], e juntou os seguintes documentos para embasar suas alegações: [LISTAR DOCUMENTOS APRESENTADOS PELO RÉU].\n\n"
        
        f"Em decisão interlocutória, o magistrado [DESCREVER A DECISÃO INTERLOCUTÓRIA, SE HOUVER], determinando [DECISÃO PROFERIDA, COMO PRODUÇÃO DE PROVAS, PERÍCIA, ETC.].\n\n"
        
        f"Em réplica, a parte autora reiterou [DESCREVER AS ALEGAÇÕES DA RÉPLICA], e trouxe novos documentos: [LISTAR NOVOS DOCUMENTOS APRESENTADOS PELA PARTE AUTORA, SE HOUVER].\n\n"
        
        f"Na audiência realizada em [DATA DA AUDIÊNCIA], foram discutidos [DESCREVER O QUE FOI DEBATIDO NA AUDIÊNCIA, COMO TENTATIVA DE CONCILIAÇÃO, INSTRUÇÃO, OITIVA DE TESTEMUNHAS, ETC.]. "
        f"As partes não chegaram a um acordo, e o processo prosseguiu.\n\n"
        
        f"Na sentença, o magistrado julgou [DESCREVER O TEOR DA SENTENÇA – EXEMPLO: procedente, improcedente, procedente em parte], com a fundamentação de que [EXPLICAR OS FUNDAMENTOS LEGAIS DA DECISÃO]. "
        f"A parte ré foi condenada a [DESCREVER A CONDENAÇÃO OU DECISÃO DA SENTENÇA].\n\n"
        
        f"Em agravo de instrumento interposto pela parte [NOME DA PARTE], argumentou-se que [DESCREVER OS PRINCIPAIS FUNDAMENTOS DO AGRAVO]. "
        f"O relator do agravo decidiu [DESCREVER A DECISÃO NO AGRAVO, SE HOUVER].\n\n"
        
        f"Em apelação, a parte [NOME DA PARTE] recorreu da sentença, alegando que [DESCREVER OS PONTOS RECORRIDOS NA APELAÇÃO]. "
        f"Em contrarrazões, a parte [NOME DA PARTE CONTRÁRIA] sustentou que [DESCREVER AS CONTRARRAZÕES].\n\n"
        
        f"Por fim, o tribunal proferiu decisão [DESCREVER SE A DECISÃO DO RECURSO FOI MONOCRÁTICA OU COLEGIADA], decidindo que [DESCREVER A DECISÃO FINAL DO RECURSO].\n\n"

        f"### Considerações Finais:\n"
        f"Caso seja necessário concluir o relatório ou fornecer um resumo geral:\n\n"
        f"[SÍNTESE IMPARCIAL DOS PONTOS PRINCIPAIS]. Exemplo: “O presente processo discute a responsabilidade da fabricante por defeitos de fabricação em produto adquirido pela parte autora. Após perícia técnica e instrução processual, o juiz de primeira instância condenou a ré à substituição do produto e ao pagamento de indenização por danos morais, decisão esta mantida pelo tribunal em sede de apelação.”\n\n"
        
        f"Texto base para análise: {text}"
    )
    
    response = model.generate_content(question)
    response.resolve()
    return response.text.strip()



# Função para criar documento DOCX modificado
def create_docx(filename, summary, docx_template_path, output_folder, initial_height_offset=0):
    output_filename = os.path.join(output_folder, os.path.splitext(filename)[0] + '_modificado.docx')
    shutil.copyfile(docx_template_path, output_filename)
    doc = Document(output_filename)
    add_text_to_second_page(doc, summary, initial_height_offset)
    doc.save(output_filename)
    return output_filename

# Função para adicionar texto à segunda página do DOCX
def add_text_to_second_page(doc, text, initial_height_offset=0):
    section = doc.sections[-1]
    new_paragraph = doc.add_paragraph()
    run = new_paragraph.add_run(text)  # Adiciona o texto sem formatação
    new_paragraph.style = doc.styles["Texto formatado"]


# Função para criar um arquivo ZIP com os arquivos fornecidos
def create_zip(files):
    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, "w") as zip_file:
        for file in files:
            zip_file.write(file, os.path.basename(file))
    zip_buffer.seek(0)
    return zip_buffer


# Função para criar um arquivo ZIP com os documentos
def create_zip(docx_file, pdf_file):
    zip_path = tempfile.mktemp(suffix=".zip")
    with zipfile.ZipFile(zip_path, 'w') as zipf:
        zipf.write(docx_file, os.path.basename(docx_file))
        zipf.write(pdf_file, os.path.basename(pdf_file))
    return zip_path


import streamlit as st
import os
import tempfile
from time import sleep

# Configura a página primeiro, antes de qualquer outra chamada do Streamlit
st.set_page_config(page_title="Gerador de Relatório Jurídico", page_icon="📄", layout="centered")

# Função principal do aplicativo Streamlit
def main():
    # Título e introdução
    st.title("📄 Gerador de Relatório Jurídico")
    st.write("**Automatize a criação de relatórios jurídicos completos e abrangentes a partir de documentos processuais.**")

    # Inicializa os estados de sessão para controlar o fluxo da aplicação
    if "uploader_key" not in st.session_state:
        st.session_state["uploader_key"] = 1
    if "uploaded_file" not in st.session_state:
        st.session_state.uploaded_file = None
    if "selected_model" not in st.session_state:
        st.session_state.selected_model = None
    if "conversion_done" not in st.session_state:
        st.session_state.conversion_done = False
    if "downloaded" not in st.session_state:
        st.session_state.downloaded = False

    # Exibe a área de upload de arquivo
    st.markdown("### Faça upload do seu documento PDF")
    uploaded_file = st.file_uploader(
        "Faça upload de um arquivo PDF",
        type="pdf",
        accept_multiple_files=False,
        key=str(st.session_state["uploader_key"]),
    )
    
    if uploaded_file and not st.session_state.conversion_done:
        st.session_state.uploaded_file = uploaded_file

    # Só exibe o selectbox de modelo se o arquivo foi enviado
    if st.session_state.uploaded_file and not st.session_state.conversion_done:
        # Menu para escolher o modelo de documento
        st.markdown("### Escolha o modelo para o relatório:")
        model_choice = st.selectbox(
            "", 
            [
                "Selecione um modelo",
                "Urban Pulse – Design empresarial dinâmico e moderno",
                "Emerald Horizon - Design sofisticado com toques modernos em verde-menta",
                "Visionary Blueprint - Sofisticado e futurista, perfeito para relatórios estratégicos",
                "Corporate Pulse – Elegante e refinado, transmitindo confiança e organização",
                "Mint Elegance - Um toque elegante e minimalista, formal na medida certa"
            ]
        )
        
        # Define o caminho do modelo de acordo com a escolha
        model_paths = {
            "Urban Pulse – Design empresarial dinâmico e moderno": 'Canva/Modelo1.docx',
            "Emerald Horizon - Design sofisticado com toques modernos em verde-menta": 'Canva/Modelo2.docx',
            "Visionary Blueprint - Sofisticado e futurista, perfeito para relatórios estratégicos": 'Canva/Modelo3.docx',
            "Corporate Pulse – Elegante e refinado, transmitindo confiança e organização": 'Canva/Modelo4.docx',
            "Mint Elegance - Um toque elegante e minimalista, formal na medida certa": 'Canva/Modelo5.docx'
        }

        if model_choice in model_paths:
            st.session_state.selected_model = model_paths[model_choice]
        else:
            st.session_state.selected_model = None

        # Exibe o botão de "Resumir documento" se o modelo foi selecionado
        if st.session_state.selected_model and st.button("Resumir documento"):
            # Processa o arquivo e gera o resumo
            with tempfile.NamedTemporaryFile(delete=False) as tmp_file:
                tmp_file.write(st.session_state.uploaded_file.getbuffer())
                pdf_path = tmp_file.name

            st.success(f"Arquivo {st.session_state.uploaded_file.name} enviado com sucesso!")

            # Simula o processamento
            with st.spinner("Processando..."):
                sleep(3)  # Simulando algum tempo de processamento

            # Extrai texto do PDF (Função que você já deve ter definida)
            text = extract_text_from_pdf(pdf_path)

            # Gera resumo usando Gemini (Função que você já deve ter definida)
            summary = generate_summary(text)

            # Diretório persistente para salvar os arquivos DOCX modificados
            output_folder = 'Resumos/Word'
            if not os.path.exists(output_folder):
                os.makedirs(output_folder)

            # Cria o DOCX modificado usando o modelo selecionado (Função que você já deve ter definida)
            docx_file = create_docx(st.session_state.uploaded_file.name, summary, st.session_state.selected_model, output_folder, initial_height_offset=5)

            # Armazena o caminho do DOCX gerado na sessão
            st.session_state.docx_file = docx_file

            # Marca o processo como concluído
            st.session_state.conversion_done = True

    # Exibe o botão de download se a conversão estiver concluída
    if st.session_state.conversion_done:
        docx_file = st.session_state.docx_file

        # Mensagens de conversão
        st.success("Conversão concluída!")

        # Botão para baixar o documento DOCX
        with open(docx_file, "rb") as docx_fp:
            if st.download_button(
                label="Baixar Relatório (DOCX)",
                data=docx_fp,
                file_name="relatorio_juridico.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ):
                st.session_state.downloaded = True

        # Exibe o botão de "Resumir outro processo" apenas após o download
        if st.session_state.downloaded:
            if st.button("Resumir outro processo"):
                # Altera a chave do uploader para limpá-lo
                st.session_state["uploader_key"] += 1

                # Limpa os estados da sessão relacionados ao arquivo e ao progresso
                st.session_state.uploaded_file = None
                st.session_state.selected_model = None
                st.session_state.conversion_done = False
                st.session_state.downloaded = False

                # Recarrega a página para refletir as mudanças
                st.rerun()  # Força o recarregamento da página

if __name__ == "__main__":
    main()
